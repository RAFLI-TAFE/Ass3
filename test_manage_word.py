# -*- coding: utf-8 -*-
"""
Created on Sat Nov  2 14:53:40 2024

@author: Rafli (20115243)
PartB Unittest
"""
import unittest # Import unittest library for creating tests
import os # Import os library to check if files exist
from docx import Document # Import Document class to handle Word documents
from manage_word import WordManager  # Import WordManager class from your program

class TestWordManager(unittest.TestCase):  # Create a test class

    def setUp(self): # Set up before each test. This runs before every test
        self.filename = "test_document.docx"  # The name of the test document
        self.word_manager = WordManager(self.filename)  # Creating an instance of WordManager
        self.image_filename = "test_image.png"  # Image filename
        
        # Make a simple image if it doesn't exist
        if not os.path.exists(self.image_filename):
            self.create_test_image()

    def tearDown(self): # Clean up after each test. This runs after every test
        if os.path.exists(self.filename):  # If the test document exists
            os.remove(self.filename)  # Delete it

    def create_test_image(self): # Create a simple image for testing
        from PIL import Image  # Importing Image from Pillow
        image = Image.new('RGB', (100, 100), color='white')  # Make a 100x100 white image
        image.save(self.image_filename)  # Save the image

    def test_create_document(self): # Test if the document is created
        self.word_manager.create_document()  # Create the document
        self.assertTrue(os.path.exists(self.filename), "Document should be created.")  # Check if it exists

    def test_add_paragraph(self): # Test if a paragraph can be added
        self.word_manager.create_document()  # Create the document
        self.word_manager.add_paragraph()  # Add a paragraph
        
        doc = Document(self.filename)  # Open the document
        # Check if there are more than 2 paragraphs
        self.assertGreater(len(doc.paragraphs), 2, "Should have more than 2 paragraphs after adding one.") # Assert that the document is created

    def test_add_picture(self): # Test if a picture can be added
        self.word_manager.create_document()  # Create the document
        self.word_manager.add_picture(self.image_filename)  # Add the picture
        
        doc = Document(self.filename)  # Open the document
        # Check if there is at least 1 picture in the document
        self.assertGreater(len(doc.inline_shapes), 0, "Should have at least 1 picture!")  # Assert that a picture is added

    def test_get_full_text(self): # Test if we can get all the text from the document
        self.word_manager.create_document()  # Create the document
        self.word_manager.add_paragraph()  # Add a paragraph
        
        full_text = self.word_manager.get_full_text()  # Get all text
        # Check if specific text is in the full text
        self.assertIn("This is a paragraph under Main heading", full_text, "Text should include the paragraph!") # Assert that the text is included

if __name__ == "__main__":  # Main function to run the tests
    unittest.main()  # Run all the tests