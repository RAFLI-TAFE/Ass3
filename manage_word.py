# -*- coding: utf-8 -*-
"""
Created on Sat Nov  1 12:03:07 2024

@author: Rafli (20115423)

Part B 
In the second part of this assessment, you must design, code, and test a Python program that uses Python to manage WORD files. 
1.	Write a Python program that provides the ability to: 
1.1.	Create a word document with at least two headings.
1.2.	Add paragraphs to headings from task 1.1.
1.3.	Add a picture to your word document.
1.4.	Get full text from word document.
"""

from docx import Document  # Import Document class for creating Word documents
from docx.shared import Inches  # Import Inches for picture resizing
import os # Importing os for file handling

class WordManager:
    
    def __init__(self, filename): #Initialize the WordManager with a filename
        self.filename = filename  # Store the filename
        self.document = Document()  # Create a new document instance

    def create_document(self): # Create a document with headings and some text
        # Add main heading and subheading
        self.document.add_heading('Main Heading', level=1)  # Add main heading
        self.document.add_heading('Subheading', level=2)  # Add subheading
        
        self.add_paragraph()  # Call to add the first paragraph
        
    def add_paragraph(self): # Add sample paragraphs to the document

        # Sample text for the paragraph under subheading
        sample_paragraph = (
            "One nice day, I went to the park. The sun was shining, and the sky was blue."
            "I sat on the grass and felt happy. I heard birds singing and saw flowers."
            "I thought about fun things to do. I wanted to go to new places and meet new people."
            "I took out my notebook and wrote my ideas down."
            "I wanted to remember this happy day and all the good things around me."
            )
        self.document.add_paragraph(sample_paragraph)  # Add sample paragraph text
        main_heading = self.document.paragraphs[1]  # Main heading is the first paragraph
        main_heading.insert_paragraph_before("This is a paragraph under Main heading")  # Insert paragraph after the main heading
        self.save_document()  # Save the document after adding the paragraph
        
    def add_picture(self, filename): # Add a picture to the document
        # Construct the full path to the image file
        image_path = os.path.join(os.getcwd(), filename)  # Get the current working directory
        if os.path.exists(image_path): # Check if the image file exists
            self.document.add_picture(image_path, width=Inches(2))  # Add picture with smaller size
            self.save_document()  # Save the document after adding picture
        else:
            print("Image not found:", image_path)  # Notify if image not found
    
    def get_full_text(self): # Get all text from the document
        text_list = []  # Create a list to hold the text
        for para in self.document.paragraphs:  # Loop through paragraphs
            text_list.append(para.text)  # Add each paragraph text to the list
        return "\n".join(text_list)  # Join the list into a single string
        
    def save_document(self): # Save the document with the filename
        self.document.save(self.filename)  # Save the document with the filename
        
        
def play():
    filename_input = input("Enter the filename for the Word document (e.g., example.docx): ")
    
    if filename_input.endswith('.docx'):
        word_manager = WordManager(filename_input)
        word_manager.create_document()  # Create document with headings and paragraphs
        
        # Optional: add picture or get full text as needed
        action1 = input("Would you like to add a picture? (yes/no) ").lower()
        if action1 == 'yes':
            image_path = input("Enter the image filename (must be in the current directory): ")
            word_manager.add_picture(image_path)  # Add picture if available
        
        action2 = input("Would you like to get the full text? (yes/no): ").lower()
        if action2 == 'yes':
            full_text = word_manager.get_full_text()
            print("Full text of the document:\n")
            print(full_text)  # Output the full text of the document
    else:
        print("Invalid filename. Please ensure it ends with .docx.")

if __name__ == "__main__":  # Entry point of the program
    play()